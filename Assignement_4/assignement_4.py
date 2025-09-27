#!/usr/bin/env python3
"""
LLM Input Processor - A command-line tool to process text from multiple sources using LLMs.
Supports: text files, URLs (HTML), CSV files, DOCX files, PDF files, and more.
"""

import argparse
import os
import sys
from pathlib import Path
import tempfile
import requests
import csv
from urllib.parse import urlparse
import mimetypes
import openai

# Try to import required libraries
try:
    import markdown
    from docx import Document
    import PyPDF2
except ImportError as e:
    print(f"Missing required library: {e}")
    sys.exit(1)

class InputProcessor:
    """Process different input sources and extract text content."""

    def __init__(self):
        self.content = ""

    def process_file(self, file_path):
        """Process a single file based on its extension."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        mime_type, _ = mimetypes.guess_type(file_path)

        if file_path.suffix.lower() in ['.txt', '.md']:
            return self._process_text_file(file_path)
        elif file_path.suffix.lower() == '.csv':
            return self._process_csv_file(file_path)
        elif file_path.suffix.lower() == '.docx':
            return self._process_docx_file(file_path)
        elif file_path.suffix.lower() in ['.pdf']:
            return self._process_pdf_file(file_path)
        elif mime_type and mime_type.startswith('text/'):
            return self._process_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

    def _process_text_file(self, file_path):
        """Process text files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _process_csv_file(self, file_path):
        """Process CSV files."""
        content = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                content.append(' '.join(row.values()))
        return '\n'.join(content)

    def _process_docx_file(self, file_path):
        """Process DOCX files."""
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def _process_pdf_file(self, file_path):
        """Process PDF files."""
        content = []
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    content.append(page.extract_text())
        except Exception as e:
            # Fallback for newer PyPDF2 versions or other issues
            try:
                import pypdf
                with open(file_path, 'rb') as f:
                    pdf_reader = pypdf.PdfReader(f)
                    for page in pdf_reader.pages:
                        content.append(page.extract_text())
            except ImportError:
                raise ImportError("PyPDF2 failed, please install pypdf: pip install pypdf")
        return '\n'.join(content)

    def process_url(self, url):
        """Process a URL (HTML page)."""
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()

            # Simple HTML to text extraction
            try:
                from bs4 import BeautifulSoup
            except ImportError:
                raise ImportError("Please install beautifulsoup4: pip install beautifulsoup4")

            soup = BeautifulSoup(response.content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text and clean it up
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)

            return text
        except Exception as e:
            raise Exception(f"Error processing URL {url}: {str(e)}")

    def add_content(self, content):
        """Add content to the processor."""
        self.content += content + "\n\n"

def call_llm_api(prompt, model):
    """Call the OpenAI API to get a response."""
    try:
        # Check if API key is set
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise EnvironmentError("OPENAI_API_KEY environment variable not set")

        openai.api_key = api_key

        response = openai.ChatCompletion.create(
            model=model,
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )

        return response.choices[0].message.content.strip()
    except Exception as e:
        raise Exception(f"Error calling LLM API: {str(e)}")

def save_result(result, output_path):
    """Save the result to the specified output file."""
    output_path = Path(output_path)

    if output_path.suffix.lower() == '.docx':
        # Save as DOCX
        from docx import Document
        doc = Document()
        doc.add_paragraph(result)
        doc.save(output_path)
    else:
        # Save as text file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result)

def main():
    parser = argparse.ArgumentParser(
        description="LLM Input Processor - Process various input sources using LLMs",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s input.txt
  %(prog)s file1.txt file2.pdf --query "Explain the key points"
  %(prog)s https://example.com/page.html   --output result.txt
  %(prog)s data.csv docx_file.docx --query "Summarize the content"
        """
    )

    parser.add_argument(
        'inputs',
        nargs='+',
        help='Input sources (files or URLs)'
    )

    parser.add_argument(
        '--query',
        '-q',
        default='Summarize the following text:',
        help='Query prompt for LLM (default: "Summarize the following text:")'
    )

    parser.add_argument(
        '--output',
        '-o',
        help='Output file path (default: stdout)'
    )

    parser.add_argument(
        '--model',
        '-m',
        default='gpt-3.5-turbo',
        help='LLM model to use (default: gpt-3.5-turbo)'
    )

    args = parser.parse_args()

    # Validate inputs
    if not args.inputs:
        parser.print_help()
        sys.exit(1)

    try:
        processor = InputProcessor()

        # Process all inputs
        for input_source in args.inputs:
            # Check if it's a URL
            parsed_url = urlparse(input_source)
            if parsed_url.scheme in ('http', 'https'):
                print(f"Processing URL: {input_source}")
                content = processor.process_url(input_source)
            else:
                print(f"Processing file: {input_source}")
                content = processor.process_file(input_source)

            processor.add_content(content)

        # Get the final content
        full_content = processor.content.strip()

        if not full_content:
            print("No content to process")
            sys.exit(1)

        # Format the prompt
        prompt = f"{args.query}\n\n{full_content}"

        print("=== Input Content ===")
        print(full_content[:500] + ("..." if len(full_content) > 500 else ""))
        print("\n=== Query Prompt ===")
        print(prompt)

        # Call the LLM API
        print(f"\nCalling {args.model} API...")
        response = call_llm_api(prompt, args.model)

        # Output result
        if args.output:
            save_result(response, args.output)
            print(f"\nResult written to: {args.output}")
        else:
            print("\n=== LLM Response ===")
            print(response)

    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == '__main__':
    main()
